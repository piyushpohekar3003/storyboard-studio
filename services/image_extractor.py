"""
Extract text and images from DOCX, PDF, and Google Docs.
Each extraction function returns (text: str, images: list[dict]).
"""

import os
import io
import base64
from PIL import Image

from services.doc_fetcher import fetch_google_doc, is_google_doc_url


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _image_dims(raw_bytes):
    """Return (width, height) for image bytes, or (0, 0) on failure."""
    try:
        with Image.open(io.BytesIO(raw_bytes)) as img:
            return img.size  # (width, height)
    except Exception:
        return (0, 0)


def _ext_to_media_type(ext):
    """Map a file extension to an image media type."""
    ext = ext.lower().lstrip(".")
    mapping = {
        "png": "image/png",
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "gif": "image/gif",
        "bmp": "image/bmp",
        "tiff": "image/tiff",
        "tif": "image/tiff",
        "webp": "image/webp",
        "svg": "image/svg+xml",
    }
    return mapping.get(ext, "image/png")


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def extract_from_docx(file_path):
    """Extract text and images from a DOCX file.

    Returns:
        (text, images) — text is the full document text; images is a list of
        dicts with keys media_type, data (base64), description, width, height.
    """
    from docx import Document

    doc = Document(file_path)

    # --- Text ---
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    table_rows = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(" | ".join(cells))

    text = "\n".join(paragraphs)
    if table_rows:
        text += "\n\n" + "\n".join(table_rows)

    # --- Images ---
    images = []
    idx = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                image_part = rel.target_part
                raw = image_part.blob
                ext = os.path.splitext(image_part.partname)[1]
                media_type = _ext_to_media_type(ext)
                b64 = base64.b64encode(raw).decode("utf-8")
                w, h = _image_dims(raw)
                images.append({
                    "media_type": media_type,
                    "data": b64,
                    "description": f"Image {idx}: {w}x{h}px from DOCX",
                    "width": w,
                    "height": h,
                })
                idx += 1
            except Exception:
                continue

    return text, images


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def extract_from_pdf(file_path):
    """Extract text and images from a PDF file.

    Returns:
        (text, images) — same format as extract_from_docx.
    """
    import fitz  # PyMuPDF

    doc = fitz.open(file_path)

    pages_text = []
    images = []
    idx = 0

    for page_num, page in enumerate(doc, start=1):
        pages_text.append(page.get_text())

        for img_info in page.get_images(full=True):
            xref = img_info[0]
            try:
                base_image = doc.extract_image(xref)
                raw = base_image["image"]
                ext = base_image.get("ext", "png")
                media_type = _ext_to_media_type(ext)
                b64 = base64.b64encode(raw).decode("utf-8")
                w, h = _image_dims(raw)
                images.append({
                    "media_type": media_type,
                    "data": b64,
                    "description": f"Image {idx}: {w}x{h}px from page {page_num}",
                    "width": w,
                    "height": h,
                })
                idx += 1
            except Exception:
                continue

    doc.close()
    text = "\n".join(pages_text)
    return text, images


# ---------------------------------------------------------------------------
# Google Docs
# ---------------------------------------------------------------------------

def extract_from_gdoc(url):
    """Extract text and images from a Google Doc URL.

    Wraps the existing fetch_google_doc() and enriches each image with
    width, height, and a description string.

    Returns:
        (text, images) — same format as the other extractors.
    """
    text, raw_images = fetch_google_doc(url)

    images = []
    for idx, img in enumerate(raw_images):
        raw = base64.b64decode(img["data"])
        w, h = _image_dims(raw)
        images.append({
            "media_type": img.get("media_type", "image/png"),
            "data": img["data"],
            "description": f"Image {idx}: {w}x{h}px from Google Doc",
            "width": w,
            "height": h,
        })

    return text, images


# ---------------------------------------------------------------------------
# Auto-detect dispatcher
# ---------------------------------------------------------------------------

def extract_from_file(file_path_or_url):
    """Auto-detect format and extract text + images.

    Supports:
        - .docx files
        - .pdf files
        - Google Docs URLs

    Returns:
        (text, images)

    Raises:
        ValueError if the format is not recognised.
    """
    value = file_path_or_url.strip()

    # Google Docs URL
    if is_google_doc_url(value):
        return extract_from_gdoc(value)

    # Local file — check extension
    ext = os.path.splitext(value)[1].lower()

    if ext == ".docx":
        return extract_from_docx(value)
    elif ext == ".pdf":
        return extract_from_pdf(value)
    else:
        raise ValueError(
            f"Unsupported file format: '{ext}'. "
            "Supported formats: .docx, .pdf, or a Google Docs URL."
        )


# ---------------------------------------------------------------------------
# Save images to disk
# ---------------------------------------------------------------------------

def save_images_to_disk(images, output_dir):
    """Save extracted images to disk as img_0.png, img_1.png, etc.

    Args:
        images: list of image dicts (must have 'data' key with base64 string).
        output_dir: directory to save images into (created if needed).

    Returns:
        List of absolute file paths for the saved images.
    """
    os.makedirs(output_dir, exist_ok=True)

    saved_paths = []
    for i, img in enumerate(images):
        raw = base64.b64decode(img["data"])
        file_path = os.path.join(output_dir, f"img_{i}.png")

        # Convert to PNG via PIL to ensure consistent format
        try:
            pil_img = Image.open(io.BytesIO(raw))
            pil_img.save(file_path, format="PNG")
        except Exception:
            # Fallback: write raw bytes directly
            with open(file_path, "wb") as f:
                f.write(raw)

        saved_paths.append(os.path.abspath(file_path))

    return saved_paths
