import re
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT


def markdown_to_docx(content, title="Storyboard"):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    doc.add_heading(title, level=0)

    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # Heading
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)

        # Table
        elif line.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1  # back up one since the outer loop will increment

            # Parse table
            if len(table_lines) >= 2:
                headers = [c.strip() for c in table_lines[0].split("|")[1:-1]]
                data_rows = []
                for tl in table_lines[2:]:  # skip separator row
                    cells = [c.strip() for c in tl.split("|")[1:-1]]
                    if cells:
                        data_rows.append(cells)

                num_cols = len(headers)
                table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Headers
                for j, h in enumerate(headers):
                    cell = table.rows[0].cells[j]
                    cell.text = h
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(10)

                # Data
                for r_idx, row_data in enumerate(data_rows):
                    for c_idx, cell_text in enumerate(row_data):
                        if c_idx < num_cols:
                            table.rows[r_idx + 1].cells[c_idx].text = cell_text

                # Set column widths for 2-column storyboard
                if num_cols == 2:
                    for row in table.rows:
                        row.cells[0].width = Inches(3.5)
                        row.cells[1].width = Inches(3.5)

        # Regular paragraph
        elif line:
            p = doc.add_paragraph()
            # Handle bold
            parts = re.split(r"(\*\*.*?\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        i += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def bytes_visuals_to_pdf(visuals_json, script_text="", title="A1 Bytes — Visual Storyboard"):
    """Export visual storyboard as a PDF with real images matching the web frame strip."""
    import os
    import requests as req
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    buffer = io.BytesIO()
    page_w, page_h = landscape(A4)
    c = canvas.Canvas(buffer, pagesize=landscape(A4))

    # Paths
    base_dir = os.path.dirname(os.path.dirname(__file__))
    aparna_path = os.path.join(base_dir, 'static', 'aparna.png')
    has_aparna = os.path.exists(aparna_path)

    # B-roll image map (same as frames.js)
    BROLL_URLS = {
        'default': 'https://images.unsplash.com/photo-1611974789855-9c2a0a7236a3?w=400&h=700&fit=crop',
        'electric': 'https://images.unsplash.com/photo-1593941707882-a5bba14938c7?w=400&h=700&fit=crop',
        'motorcycle': 'https://images.unsplash.com/photo-1558618666-fcd25c85f82e?w=400&h=700&fit=crop',
        'battery': 'https://images.unsplash.com/photo-1619642751034-765dfdf7c58e?w=400&h=700&fit=crop',
        'power': 'https://images.unsplash.com/photo-1473341304170-971dccb5ac1e?w=400&h=700&fit=crop',
        'car': 'https://images.unsplash.com/photo-1492144534655-ae79c964c9d7?w=400&h=700&fit=crop',
        'solar': 'https://images.unsplash.com/photo-1509391366360-2e959784a276?w=400&h=700&fit=crop',
        'office': 'https://images.unsplash.com/photo-1497366216548-37526070297c?w=400&h=700&fit=crop',
        'server': 'https://images.unsplash.com/photo-1558494949-ef010cbdcc31?w=400&h=700&fit=crop',
        'handshake': 'https://images.unsplash.com/photo-1521791136064-7986c2920216?w=400&h=700&fit=crop',
        'money': 'https://images.unsplash.com/photo-1554672408-730436b60dde?w=400&h=700&fit=crop',
        'cash': 'https://images.unsplash.com/photo-1554672408-730436b60dde?w=400&h=700&fit=crop',
        'grocery': 'https://images.unsplash.com/photo-1542838132-92c53300491e?w=400&h=700&fit=crop',
        'coal': 'https://images.unsplash.com/photo-1611273426858-450d8e3c9fce?w=400&h=700&fit=crop',
        'retail': 'https://images.unsplash.com/photo-1441986300917-64674bd600d8?w=400&h=700&fit=crop',
        'product': 'https://images.unsplash.com/photo-1556228578-0d85b1a4d571?w=400&h=700&fit=crop',
        'stock': 'https://images.unsplash.com/photo-1611974789855-9c2a0a7236a3?w=400&h=700&fit=crop',
        'grid': 'https://images.unsplash.com/photo-1473341304170-971dccb5ac1e?w=400&h=700&fit=crop',
        'ola': 'https://images.unsplash.com/photo-1593941707882-a5bba14938c7?w=400&h=700&fit=crop',
        'tcs': 'https://images.unsplash.com/photo-1497366216548-37526070297c?w=400&h=700&fit=crop',
        'bosch': 'https://images.unsplash.com/photo-1492144534655-ae79c964c9d7?w=400&h=700&fit=crop',
        'tata': 'https://images.unsplash.com/photo-1473341304170-971dccb5ac1e?w=400&h=700&fit=crop',
        'honasa': 'https://images.unsplash.com/photo-1556228578-0d85b1a4d571?w=400&h=700&fit=crop',
    }

    # Image cache
    _img_cache = {}

    def _get_image(url):
        if url in _img_cache:
            return _img_cache[url]
        try:
            resp = req.get(url, timeout=8)
            if resp.status_code == 200:
                img = ImageReader(io.BytesIO(resp.content))
                _img_cache[url] = img
                return img
        except Exception:
            pass
        return None

    def _get_broll_url(desc):
        if not desc:
            return BROLL_URLS['default']
        d = desc.lower()
        for k, v in BROLL_URLS.items():
            if k in d:
                return v
        return BROLL_URLS['default']

    def _get_stock_url(supers):
        if not supers:
            return BROLL_URLS['stock']
        name = supers[0].lower()
        for k, v in BROLL_URLS.items():
            if k in name:
                return v
        return BROLL_URLS['stock']

    # Frame dimensions
    FW = 140  # frame width in points
    FH = FW * 16 / 9  # 9:16 ratio
    MARGIN = 10 * mm
    GAP = 12
    SHOTS_PER_ROW = 4
    TEXT_AREA = 65  # space below frame for text

    type_labels = {'presenter': 'Presenter', 'broll': 'B-Roll', 'presenter_stockcard': 'Stock Card'}
    border_colors = {
        'presenter': (0.55, 0.72, 0.95),
        'broll': (0.45, 0.82, 0.55),
        'presenter_stockcard': (0.95, 0.75, 0.3),
    }

    def _draw_pill(c, text, cx, cy, font_size=7):
        tw = c.stringWidth(text[:25], 'Helvetica-Bold', font_size)
        pw = max(tw + 12, 35)
        ph = font_size + 7
        c.saveState()
        c.setFillColor(colors.Color(0.29, 0.44, 0.65, 0.92))
        c.roundRect(cx - pw/2, cy - ph/2, pw, ph, ph/2, fill=1, stroke=0)
        c.setFillColor(colors.white)
        c.setFont('Helvetica-Bold', font_size)
        c.drawCentredString(cx, cy - font_size * 0.35, text[:25])
        c.restoreState()

    def _draw_subtitle(c, text, cx, cy, font_size=6.5):
        tw = c.stringWidth(text[:40], 'Helvetica', font_size)
        bw = tw + 14
        bh = font_size + 8
        c.saveState()
        c.setFillColor(colors.Color(0, 0, 0, 0.55))
        c.roundRect(cx - bw/2, cy - bh/2, bw, bh, 3, fill=1, stroke=0)
        c.setFillColor(colors.white)
        c.setFont('Helvetica', font_size)
        c.drawCentredString(cx, cy - font_size * 0.35, text[:40])
        c.restoreState()

    def _draw_shot_label(c, x, y, shot_num, duration):
        c.saveState()
        c.setFillColor(colors.Color(0, 0, 0, 0.6))
        c.roundRect(x + 3, y + 3, 55, 12, 6, fill=1, stroke=0)
        c.setFillColor(colors.white)
        c.setFont('Helvetica', 5.5)
        c.drawCentredString(x + 30.5, y + 6, f"SHOT {shot_num} · {duration}")
        c.restoreState()

    def draw_presenter_frame(c, x, y, shot):
        """Draw presenter frame with aparna.png at (x, y) bottom-left."""
        supers = shot.get('supers', [])
        vtype = shot.get('visual_type', 'presenter')
        br, bg, bb = border_colors.get(vtype, (0.5, 0.5, 0.5))

        # Beige background
        c.saveState()
        c.setFillColor(colors.Color(0.91, 0.88, 0.83))
        c.setStrokeColor(colors.Color(br, bg, bb))
        c.setLineWidth(2)
        c.roundRect(x, y, FW, FH, 6, fill=1, stroke=1)

        # Clip to frame and draw aparna
        c.saveState()
        p = c.beginPath()
        p.roundRect(x, y, FW, FH, 6)
        c.clipPath(p, stroke=0)
        if has_aparna:
            # Draw image covering full frame, anchored from top
            img_h = FH
            img_w = FW
            c.drawImage(aparna_path, x, y, width=img_w, height=img_h,
                        preserveAspectRatio=True, anchor='n', mask='auto')
        c.restoreState()

        # Bottom vignette for readability
        for i in range(20):
            alpha = i / 20.0 * 0.5
            c.setFillColor(colors.Color(0, 0, 0, alpha))
            c.rect(x, y + (20 - i) * (FH * 0.4 / 20), FW, FH * 0.4 / 20, fill=1, stroke=0)

        # Supers
        for i, s in enumerate(supers[:4]):
            _draw_pill(c, s, x + FW/2, y + FH * 0.32 - i * 14)

        # Subtitle only if no supers
        if not supers and shot.get('voiceover'):
            _draw_subtitle(c, shot['voiceover'][:45], x + FW/2, y + FH * 0.25)

        # Shot label
        _draw_shot_label(c, x, y, shot.get('shot', ''), shot.get('duration_hint', ''))
        c.restoreState()

    def draw_broll_frame(c, x, y, shot):
        """Draw b-roll frame with stock footage at (x, y) bottom-left."""
        supers = shot.get('supers', [])
        desc = shot.get('broll_description', '')
        br, bg, bb = border_colors.get('broll', (0.45, 0.82, 0.55))

        c.saveState()
        # Dark background
        c.setFillColor(colors.Color(0.14, 0.16, 0.20))
        c.setStrokeColor(colors.Color(br, bg, bb))
        c.setLineWidth(2)
        c.roundRect(x, y, FW, FH, 6, fill=1, stroke=1)

        # Clip and draw stock image
        c.saveState()
        p = c.beginPath()
        p.roundRect(x, y, FW, FH, 6)
        c.clipPath(p, stroke=0)
        url = _get_broll_url(desc)
        img = _get_image(url)
        if img:
            c.drawImage(img, x, y, width=FW, height=FH, preserveAspectRatio=True, anchor='c', mask='auto')
        c.restoreState()

        # Vignette overlays (top + bottom)
        for i in range(15):
            alpha = (15 - i) / 15.0 * 0.3
            c.setFillColor(colors.Color(0, 0, 0, alpha))
            c.rect(x, y + FH - (i+1) * (FH*0.2/15), FW, FH*0.2/15, fill=1, stroke=0)
        for i in range(20):
            alpha = i / 20.0 * 0.5
            c.setFillColor(colors.Color(0, 0, 0, alpha))
            c.rect(x, y + (20 - i) * (FH * 0.4 / 20), FW, FH * 0.4 / 20, fill=1, stroke=0)

        # B-roll description label (top)
        if desc:
            c.setFillColor(colors.Color(0, 0, 0, 0.5))
            tw = min(len(desc) * 4 + 10, FW - 10)
            c.roundRect(x + 5, y + FH - 18, tw, 14, 4, fill=1, stroke=0)
            c.setFillColor(colors.Color(1, 1, 1, 0.8))
            c.setFont('Helvetica', 5.5)
            c.drawString(x + 9, y + FH - 14, desc[:35])

        # Supers
        for i, s in enumerate(supers[:4]):
            _draw_pill(c, s, x + FW/2, y + FH * 0.30 - i * 14)

        if not supers and shot.get('voiceover'):
            _draw_subtitle(c, shot['voiceover'][:45], x + FW/2, y + FH * 0.25)

        _draw_shot_label(c, x, y, shot.get('shot', ''), shot.get('duration_hint', ''))
        c.restoreState()

    def draw_stockcard_frame(c, x, y, shot):
        """Draw stock card frame with stock footage + price overlay."""
        supers = shot.get('supers', [])
        stock_name = supers[0] if supers else 'STOCK'
        price_change = supers[1] if len(supers) > 1 else ''
        price_val = supers[2] if len(supers) > 2 else ''
        is_up = '↑' in price_change or '+' in price_change
        chg_color = colors.Color(0.13, 0.77, 0.37) if is_up else colors.Color(0.94, 0.27, 0.27)
        br, bg, bb = border_colors.get('presenter_stockcard', (0.95, 0.75, 0.3))

        c.saveState()
        c.setFillColor(colors.Color(0.12, 0.14, 0.18))
        c.setStrokeColor(colors.Color(br, bg, bb))
        c.setLineWidth(2)
        c.roundRect(x, y, FW, FH, 6, fill=1, stroke=1)

        # Clip and draw stock footage
        c.saveState()
        p = c.beginPath()
        p.roundRect(x, y, FW, FH, 6)
        c.clipPath(p, stroke=0)
        url = _get_stock_url(supers)
        img = _get_image(url)
        if img:
            c.drawImage(img, x, y, width=FW, height=FH, preserveAspectRatio=True, anchor='c', mask='auto')
        c.restoreState()

        # Dark overlay for the top price card area
        c.setFillColor(colors.Color(0, 0, 0, 0.35))
        c.rect(x, y + FH * 0.6, FW, FH * 0.4, fill=1, stroke=0)

        # White price card
        card_x = x + 8
        card_y = y + FH * 0.65
        card_w = FW - 16
        card_h = FH * 0.28
        c.setFillColor(colors.Color(1, 1, 1, 0.95))
        c.roundRect(card_x, card_y, card_w, card_h, 5, fill=1, stroke=0)

        # Stock name
        c.setFillColor(colors.Color(0.2, 0.2, 0.2))
        c.setFont('Helvetica-Bold', 8)
        c.drawString(card_x + 8, card_y + card_h - 14, stock_name[:20])

        # Price value
        c.setFillColor(colors.Color(0.1, 0.1, 0.1))
        c.setFont('Helvetica-Bold', 12)
        c.drawString(card_x + 8, card_y + card_h - 30, price_val[:15])

        # Change badge
        if price_change:
            c.setFillColor(colors.Color(chg_color.red, chg_color.green, chg_color.blue, 0.15))
            badge_w = c.stringWidth(price_change[:15], 'Helvetica-Bold', 7) + 10
            c.roundRect(card_x + 8, card_y + card_h - 45, badge_w, 12, 3, fill=1, stroke=0)
            c.setFillColor(chg_color)
            c.setFont('Helvetica-Bold', 7)
            c.drawString(card_x + 13, card_y + card_h - 42, price_change[:15])

        # AngelOne label
        c.setFillColor(colors.Color(1, 1, 1, 0.6))
        c.setFont('Helvetica-Bold', 6)
        c.drawRightString(x + FW - 8, y + FH - 14, 'AngelOne')

        _draw_shot_label(c, x, y, shot.get('shot', ''), shot.get('duration_hint', ''))
        c.restoreState()

    def draw_frame(c, x, y, shot):
        vtype = shot.get('visual_type', 'presenter')
        if vtype == 'presenter_stockcard':
            draw_stockcard_frame(c, x, y, shot)
        elif vtype == 'broll':
            draw_broll_frame(c, x, y, shot)
        else:
            draw_presenter_frame(c, x, y, shot)

    def draw_text_below(c, x, y, shot, max_w):
        """Draw shot info text below the frame."""
        vtype = shot.get('visual_type', 'presenter')
        supers = shot.get('supers', [])
        vo = shot.get('voiceover', '')
        broll = shot.get('broll_description', '')

        ty = y  # current text y position

        # Shot label + type
        c.setFont('Helvetica-Bold', 8)
        c.setFillColor(colors.Color(0.1, 0.1, 0.1))
        label = f"Shot {shot.get('shot', '?')} · {shot.get('duration_hint', '')}"
        c.drawString(x, ty, label)
        lw = c.stringWidth(label, 'Helvetica-Bold', 8)
        c.setFont('Helvetica', 6.5)
        c.setFillColor(colors.grey)
        c.drawString(x + lw + 6, ty + 1, type_labels.get(vtype, ''))
        ty -= 11

        # Voiceover (truncated)
        c.setFont('Helvetica', 6.5)
        c.setFillColor(colors.Color(0.35, 0.35, 0.35))
        vo_display = vo[:80] + ('...' if len(vo) > 80 else '')
        # Word wrap
        words = vo_display.split()
        line = ''
        for w in words:
            test = (line + ' ' + w).strip()
            if c.stringWidth(test, 'Helvetica', 6.5) > max_w:
                c.drawString(x, ty, line)
                ty -= 9
                line = w
            else:
                line = test
        if line:
            c.drawString(x, ty, line)
            ty -= 11

        # Supers
        sx = x
        c.setFont('Helvetica-Bold', 6)
        for s in supers[:4]:
            c.setFillColor(colors.Color(0.29, 0.44, 0.65))
            tw = c.stringWidth(s[:20], 'Helvetica-Bold', 6) + 8
            if sx + tw > x + max_w:
                sx = x
                ty -= 10
            c.roundRect(sx, ty - 2, tw, 10, 3, fill=0, stroke=1)
            c.drawString(sx + 4, ty, s[:20])
            sx += tw + 4

        if supers:
            ty -= 12

        # B-roll description
        if broll:
            c.setFont('Helvetica-Oblique', 5.5)
            c.setFillColor(colors.Color(0.5, 0.5, 0.5))
            c.drawString(x, ty, f"B-roll: {broll[:40]}")

    # ── Page rendering ──

    def draw_page_header(c, page_num):
        """Draw title and header on each page."""
        if page_num == 1:
            c.setFont('Helvetica-Bold', 16)
            c.setFillColor(colors.Color(0.1, 0.1, 0.1))
            c.drawString(MARGIN, page_h - MARGIN - 5, title)

            c.setFont('Helvetica', 8)
            c.setFillColor(colors.grey)
            c.drawString(MARGIN, page_h - MARGIN - 18, f"{len(visuals_json)} shots · Angel One Bytes")

            # Script
            if script_text:
                c.setFont('Helvetica-Bold', 10)
                c.setFillColor(colors.Color(0.1, 0.1, 0.1))
                c.drawString(MARGIN, page_h - MARGIN - 38, "Script")
                c.setFont('Helvetica', 7)
                c.setFillColor(colors.Color(0.35, 0.35, 0.35))
                display = script_text[:300] + ('...' if len(script_text) > 300 else '')
                # Simple word wrap
                words = display.split()
                line = ''
                ly = page_h - MARGIN - 52
                for w in words:
                    test = (line + ' ' + w).strip()
                    if c.stringWidth(test, 'Helvetica', 7) > page_w - 2 * MARGIN:
                        c.drawString(MARGIN, ly, line)
                        ly -= 10
                        line = w
                    else:
                        line = test
                if line:
                    c.drawString(MARGIN, ly, line)
                    ly -= 10
                return ly - 10
            return page_h - MARGIN - 40
        else:
            c.setFont('Helvetica', 7)
            c.setFillColor(colors.grey)
            c.drawString(MARGIN, page_h - MARGIN - 5, f"{title} — Page {page_num}")
            return page_h - MARGIN - 18

    # Calculate layout
    total_col_w = (page_w - 2 * MARGIN - (SHOTS_PER_ROW - 1) * GAP) / SHOTS_PER_ROW
    FW_actual = min(FW, total_col_w)
    FH_actual = FW_actual * 16 / 9

    page_num = 1
    start_y = draw_page_header(c, page_num)
    cursor_y = start_y

    for row_start in range(0, len(visuals_json), SHOTS_PER_ROW):
        row_shots = visuals_json[row_start:row_start + SHOTS_PER_ROW]
        row_height = FH_actual + TEXT_AREA + 10

        # Check if row fits on current page
        if cursor_y - row_height < MARGIN:
            c.showPage()
            page_num += 1
            cursor_y = draw_page_header(c, page_num)

        # Draw frames
        for idx, shot in enumerate(row_shots):
            fx = MARGIN + idx * (total_col_w + GAP) + (total_col_w - FW_actual) / 2
            fy = cursor_y - FH_actual

            # Temporarily adjust FW/FH for drawing
            old_fw, old_fh = FW, FH
            FW, FH = FW_actual, FH_actual
            draw_frame(c, fx, fy, shot)
            FW, FH = old_fw, old_fh

            # Text below frame
            tx = MARGIN + idx * (total_col_w + GAP)
            draw_text_below(c, tx, fy - 12, shot, total_col_w)

        cursor_y -= row_height

    c.save()
    buffer.seek(0)
    return buffer


def bytes_visuals_to_docx(visuals_json, script_text="", title="A1 Bytes — Visual Storyboard"):
    """Export visual storyboard as a DOCX with a shot table."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    doc.add_heading(title, level=0)

    if script_text:
        doc.add_heading("Script", level=2)
        p = doc.add_paragraph(script_text)
        p.style.font.size = Pt(9)
        doc.add_paragraph()

    doc.add_heading("Shot-by-Shot Visual Breakdown", level=2)

    # Create table: Shot | Type | Voiceover | Supers | B-roll | Duration | Notes
    headers = ["Shot", "Type", "Voiceover", "Supers", "B-roll Guidance", "Duration", "Notes"]
    table = doc.add_table(rows=1 + len(visuals_json), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        # Shade header
        from docx.oxml.ns import qn
        shading = cell._element.get_or_add_tcPr()
        shading_el = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '4a6fa5',
            qn('w:val'): 'clear',
        })
        shading.append(shading_el)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    type_labels = {
        'presenter': 'Presenter',
        'broll': 'B-Roll',
        'presenter_stockcard': 'Stock Card',
    }

    for idx, shot in enumerate(visuals_json):
        row = table.rows[idx + 1]
        row.cells[0].text = str(shot.get('shot', idx + 1))
        row.cells[1].text = type_labels.get(shot.get('visual_type', ''), shot.get('visual_type', ''))
        row.cells[2].text = shot.get('voiceover', '')
        row.cells[3].text = '\n'.join(shot.get('supers', []))
        row.cells[4].text = shot.get('broll_description', '') or '—'
        row.cells[5].text = shot.get('duration_hint', '')
        row.cells[6].text = shot.get('notes', '') or ''

        # Style cells
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

        # Color-code the type cell
        type_colors = {
            'presenter': 'dbeafe',
            'broll': 'dcfce7',
            'presenter_stockcard': 'fef3c7',
        }
        vtype = shot.get('visual_type', '')
        if vtype in type_colors:
            tc = row.cells[1]._element.get_or_add_tcPr()
            shading_el = tc.makeelement(qn('w:shd'), {
                qn('w:fill'): type_colors[vtype],
                qn('w:val'): 'clear',
            })
            tc.append(shading_el)

    # Set column widths
    col_widths = [0.5, 0.9, 2.5, 1.8, 1.5, 0.6, 1.2]
    for row in table.rows:
        for i, w in enumerate(col_widths):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def script_to_docx(content, title="Script"):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    doc.add_heading(title, level=0)

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        else:
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*.*?\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --- Shorts Visual Storyboard PDF ---


def shorts_storyboard_to_pdf(storyboard_json, project_dir, title="Shorts Storyboard"):
    """
    Generate a PDF from the shorts visual storyboard JSON.
    Uses reportlab to draw simplified frame previews with text info.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.colors import Color
    from reportlab.pdfgen import canvas as pdf_canvas
    import os

    W, H = A4
    buffer = io.BytesIO()
    c = pdf_canvas.Canvas(buffer, pagesize=A4)

    BG = Color(6/255, 6/255, 9/255)
    GOLD = Color(255/255, 215/255, 0/255)
    CYAN = Color(0/255, 229/255, 255/255)
    MINT = Color(0/255, 229/255, 160/255)
    PURPLE = Color(179/255, 136/255, 255/255)
    WHITE = Color(240/255, 240/255, 245/255)
    MUTED = Color(106/255, 106/255, 122/255)
    BODY = Color(176/255, 176/255, 184/255)

    TYPE_COLORS = {
        "presenter": CYAN, "presenter_cards": CYAN,
        "data": GOLD, "section_header": GOLD, "cta": PURPLE,
    }

    def new_page():
        c.showPage()
        c.setFillColor(BG)
        c.rect(0, 0, W, H, fill=1, stroke=0)

    c.setFillColor(BG)
    c.rect(0, 0, W, H, fill=1, stroke=0)
    c.setFillColor(GOLD)
    c.setFont("Helvetica-Bold", 22)
    c.drawCentredString(W/2, H - 50, title)
    c.setFillColor(MUTED)
    c.setFont("Helvetica", 10)
    total = storyboard_json.get("metadata", {}).get("total_frames", 0)
    c.drawCentredString(W/2, H - 70, f"Visual Storyboard | Short (9:16) | {total} Frames")

    y = H - 110

    for section in storyboard_json.get("sections", []):
        label = section.get("section_label", "")

        if y < 120:
            new_page()
            y = H - 40
        c.setStrokeColor(Color(1, 0.84, 0, 0.2))
        c.setLineWidth(0.5)
        c.line(50, y, W - 50, y)
        c.setFillColor(BG)
        tw = c.stringWidth(label, "Helvetica-Bold", 10)
        c.rect(W/2 - tw/2 - 8, y - 5, tw + 16, 10, fill=1, stroke=0)
        c.setFillColor(GOLD)
        c.setFont("Helvetica-Bold", 10)
        c.drawCentredString(W/2, y - 3, label)
        y -= 28

        for frame in section.get("frames", []):
            if y < 300:
                new_page()
                y = H - 40

            fnum = frame.get("frame_number", 0)
            ftype = frame.get("frame_type", "presenter")
            dialogue = frame.get("dialogue", "")
            visual = frame.get("visual_cue", "")

            thumb_w, thumb_h = 120, 213
            thumb_x = 40
            thumb_y = y - thumb_h

            c.setFillColor(Color(10/255, 10/255, 15/255))
            c.setStrokeColor(Color(1, 1, 1, 0.06))
            c.setLineWidth(0.5)
            c.roundRect(thumb_x, thumb_y, thumb_w, thumb_h, 5, fill=1, stroke=1)

            svg_path = os.path.join(project_dir, f"frame-{fnum:02d}.svg")
            try:
                from svglib.svglib import svg2rlg
                from reportlab.graphics import renderPDF
                if os.path.exists(svg_path):
                    drawing = svg2rlg(svg_path)
                    if drawing:
                        sx = thumb_w / drawing.width
                        sy = thumb_h / drawing.height
                        scale = min(sx, sy)
                        drawing.width = thumb_w
                        drawing.height = thumb_h
                        drawing.scale(scale, scale)
                        renderPDF.draw(drawing, c, thumb_x, thumb_y)
            except Exception:
                c.setFillColor(MUTED)
                c.setFont("Helvetica", 7)
                c.drawCentredString(thumb_x + thumb_w/2, thumb_y + thumb_h/2, f"Frame {fnum}")

            info_x = thumb_x + thumb_w + 20

            c.setFillColor(MUTED)
            c.setFont("Helvetica-Bold", 11)
            c.drawString(info_x, y - 12, f"FRAME {fnum:02d}")

            tc = TYPE_COLORS.get(ftype, CYAN)
            c.setFillColor(Color(tc.red, tc.green, tc.blue, 0.1))
            c.setStrokeColor(Color(tc.red, tc.green, tc.blue, 0.3))
            c.setLineWidth(0.5)
            badge_text = ftype.upper().replace("_", " ")
            btw = c.stringWidth(badge_text, "Helvetica-Bold", 7) + 10
            c.roundRect(info_x, y - 30, btw, 13, 3, fill=1, stroke=1)
            c.setFillColor(tc)
            c.setFont("Helvetica-Bold", 7)
            c.drawString(info_x + 5, y - 27, badge_text)

            cur_y = y - 50

            c.setFillColor(CYAN)
            c.setFont("Helvetica-Bold", 8)
            c.drawString(info_x, cur_y, "DIALOGUE")
            cur_y -= 12
            c.setFillColor(BODY)
            c.setFont("Helvetica", 8)
            for line in (dialogue or "").split("\n")[:4]:
                c.drawString(info_x + 6, cur_y, line[:85])
                cur_y -= 11

            cur_y -= 6
            c.setFillColor(GOLD)
            c.setFont("Helvetica-Bold", 8)
            c.drawString(info_x, cur_y, "VISUAL CUE")
            cur_y -= 12
            c.setFillColor(BODY)
            c.setFont("Helvetica", 8)
            for line in (visual or "").split("\n")[:4]:
                c.drawString(info_x + 6, cur_y, line[:85])
                cur_y -= 11

            y -= max(thumb_h + 20, 240)

    c.save()
    buffer.seek(0)
    return buffer
